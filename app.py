from flask import Flask, render_template, request, send_file, redirect, url_for, session, send_file
import pandas as pd
import os
from docx import Document
import requests
import json
from readBackup import clauses_dict
from pocketbase import PocketBase

app = Flask(__name__)

app.secret_key = 'your_secret_key'  # Necessary for session management

app.config['DEBUG'] = True
app.config['PROPAGATE_EXCEPTIONS'] = True


# Initialize PocketBase client
pb = PocketBase('http://157.245.244.136:8090')  # Change this to your PocketBase URL

# Load the Excel workbook
file_path = 'Clause Matrix.xlsx'  # Replace this with the correct path to your file
xls = pd.ExcelFile(file_path, engine='openpyxl')

# Load the 'sort1' sheet into a DataFrame and force 'ALL PROCUREMENT TYPES' to be read as strings
df_sort1 = pd.read_excel(xls, sheet_name='sort1', dtype={'ALL PROCUREMENT TYPES': str})

# Convert 'ALL PROCUREMENT TYPES' column to avoid NaN issues
df_sort1['ALL PROCUREMENT TYPES'] = df_sort1['ALL PROCUREMENT TYPES'].fillna('').astype(str)

# Print the DataFrame to inspect the data for ALL PROCUREMENT TYPES
#print(df_sort1[['Cost', 'ALL PROCUREMENT TYPES']])  # <<--- Inserted here for debugging

# Initialize an empty dictionary to store clause IDs based on rows and columns
procurement_data = {}

# Iterate through each row in the DataFrame and populate the dictionary
for index, row in df_sort1.iterrows():
    cost_category = row['Cost']  # Assuming this column contains cost categories like "$10,000", "$25,000", etc.
    for column in df_sort1.columns[1:]:  # Skipping the first column (Cost), iterate over procurement types
        procurement_type = column
        clause_ids = row[procurement_type]  # Get clause IDs for this procurement type
        if cost_category not in procurement_data:
            procurement_data[cost_category] = {}
        procurement_data[cost_category][procurement_type] = clause_ids

# Function to get all cost thresholds that match the user input
def get_matching_thresholds(project_cost):
    """ Returns a list of cost thresholds that are less than or equal to the input project_cost """
    thresholds = []

    # Add ANY COST first since it's always included
    thresholds.append("ANY COST")
    
    # Add other thresholds based on project cost
    if project_cost > 10000:
        thresholds.append(">$10,000")
    if project_cost > 25000:
        thresholds.append(">$25,000")
    if project_cost > 100000:
        thresholds.append(">$100,000")
    if project_cost > 150000:
        thresholds.append(">$150,000")
    if project_cost > 250000:
        thresholds.append(">$250,000")
    
    return thresholds

# Route to the login page
@app.route('/')
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']

        try:
            # Attempt to authenticate the user with PocketBase
            auth_data = pb.collection('users').auth_with_password(email, password)

            # Save the user session
            session['user_id'] = auth_data.record.id
            session['username'] = auth_data.record.username
            session['email'] = email

            return redirect(url_for('get_clauses'))
        except Exception as e:
            # Log the error for detailed inspection
            print(f"Error: {str(e)}")
            if hasattr(e, 'response') and hasattr(e.response, 'json'):
                print("Detailed response from PocketBase:", e.response.json())
            return f"Login failed: {str(e)}"


    return render_template('login.html')

# Route to the sign-up page
# Route to the sign-up page
import json  # Required to print data as JSON

@app.route('/signup', methods=['GET', 'POST'])
def signup():
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']
        username = request.form['username']

        # Prepare the data to be sent
        user_data = {
            'email': email,
            'password': password,
            'passwordConfirm': password,  # Ensure you're sending this field
            'username': username
        }

        # Log the data being sent to PocketBase
        print("Sending request to PocketBase with user data:", json.dumps(user_data, indent=4))

        try:
            # Attempt to create a new user in PocketBase
            new_user = pb.collection('users').create(user_data)

            # Redirect to login upon successful signup
            return redirect(url_for('login'))
        
        except Exception as e:
            # Catch and log any exception that occurs
            print(f"Exception occurred: {e}")
            # Log detailed response if available
            response_error = e.response.json() if hasattr(e, 'response') else 'No response'
            print(f"Response error: {response_error}")
            # Return the error message to the user for more clarity
            return f"Sign-up failed: {str(e)}, Response: {response_error}"
    
    # Render the signup form if it's a GET request
    return render_template('signup.html')

UPLOAD_FOLDER = os.path.join(os.getcwd(), 'uploads')  # Save files in the 'uploads' folder
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)  # Create the folder if it doesn't exist

project_title = ""

@app.route('/get_clauses', methods=['GET', 'POST'])
def get_clauses():
    if 'user_id' not in session:
        return redirect(url_for('login'))

    procurement_types = df_sort1.columns[1:].tolist() # Get all columns except 'Cost'

    if request.method == 'POST':
        # Handle form submission
        if 'cost' not in request.form or request.form['cost'] == '' or 'title' not in request.form or request.form['title'] == '':
            return "Project title or cost is missing", 400  # Handle missing title or cost input
        
        selected_column = request.form['column']
        project_cost = float(request.form['cost'])  # Convert cost to float
        project_title = request.form['title']  # Get the project title
        user_id = session['user_id']  # Get the current user's ID from the session

        # Mapping text-based thresholds to numeric values
        cost_mapping = {
            ">$10,000": 10000,
            ">$25,000": 25000,
            ">$100,000": 100000,
            ">$150,000": 150000,
            ">$250,000": 250000,
            "ANY COST": 0  # Assume "ANY COST" applies to any value
        }

        # Initialize an empty list to store all relevant clause IDs
        clause_ids = []

        # Debugging: print project cost
        #print(f"User project cost: {project_cost}")
        
        # Iterate over each row in the DataFrame
        for index, row in df_sort1.iterrows():
            row_cost_label = row['Cost']

            # Get the numeric value corresponding to the cost threshold
            if row_cost_label in cost_mapping:
                row_cost = cost_mapping[row_cost_label]
            else:
                print(f"Skipping row index {index}: Invalid or unmapped cost value '{row_cost_label}'")
                continue

            # Debugging: print the current row cost and type
            #print(f"Row index: {index}, Row cost: {row_cost}, Type: {type(row_cost)}")
            
            # Check if the current row cost is less than or equal to the user's input
            if project_cost >= row_cost:
                #print(f"Row cost {row_cost} is <= user input cost {project_cost}")
                
                # Get clause IDs from the selected column
                selected_column_ids = row[selected_column]

                # Debugging: print clause IDs from the selected column
                #print(f"Selected column ({selected_column}) IDs: {selected_column_ids}")

                # Check if the value is not NaN and append to clause_ids
                if pd.notna(selected_column_ids):
                    clause_ids.extend([clause.strip() for clause in str(selected_column_ids).split(',')])

                # Get clause IDs from the "ALL PROCUREMENT TYPES" column
                all_procurement_ids = row['ALL PROCUREMENT TYPES']
                
                # Debugging: print clause IDs from the "ALL PROCUREMENT TYPES" column
                #print(f"ALL PROCUREMENT TYPES IDs: {all_procurement_ids}")

                # Check if the value is not NaN and append to clause_ids
                if pd.notna(all_procurement_ids):
                    clause_ids.extend([clause.strip() for clause in str(all_procurement_ids).split(',')])

        # Ensure the clause IDs list is unique and sorted
        clause_ids = list(set(clause_ids))  # Remove duplicates

        #print("Clause IDs fetched from Excel:", clause_ids)
        
        try:
            clause_ids = [float(clause_id) for clause_id in clause_ids if clause_id]
        except ValueError as e:
            return f"Error processing clause IDs: {str(e)}", 400
        
        # Debugging: Print the clause IDs after converting to float
        #print("Clause IDs after conversion to float:", clause_ids)

        document = Document()
        document.add_heading(f'Clauses Report for {project_title}', 0)

        for clause_id in clause_ids:
            # Debugging: Print clause ID being looked up
            #print(f"Looking up clause ID: {clause_id}")
            
            clause_data = clauses_dict.get(clause_id, {})
            
            if not clause_data:
                clause_data = clauses_dict.get(int(clause_id), {})  # Try as integer
            if not clause_data:
                clause_data = clauses_dict.get(str(int(clause_id)), {})  # Try as string (integer)
             
             # Debugging: Print the fetched clause data
            #print(f"Fetched clause data for ID {clause_id}: {clause_data}")
            title = clause_data.get('Title', f'Clause {clause_id} (No Title Found)')
            text = clause_data.get('Text', 'No Text Found')

            document.add_heading(title, level=1)
            document.add_paragraph(text)
            
        # Save the document to the 'uploads' directory
        file_name = f'{project_title}_Clauses_Report.docx'
        output_file_path = os.path.join(UPLOAD_FOLDER, file_name)
        document.save(output_file_path)

        print(f"Uploading file: {output_file_path}")  # Print the file being uploaded

         # Manual file upload using a POST request
        try:
            with open(output_file_path, 'rb') as file:
                form_data = {
                    'title': project_title,
                    'user': user_id
                }

                # Prepare multipart data with the file
                files = {
                    'file': ('file', file, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                }

                # Manually send the POST request with the form data and file
                response = requests.post(f"http://157.245.244.136:8090/api/collections/projects/records", data=form_data, files=files)

                if response.status_code == 200:
                    print("File uploaded successfully!")
                    new_project = response.json()
                else:
                    return f"File upload failed: {response.status_code}, {response.text}"

        except Exception as e:
            return f"Error saving project: {e}"


        # Fetch all projects for the current user
        try:
            user_projects = pb.collection('projects').get_full_list(query_params={
                'filter': f'user = "{user_id}"'
            })

            # Generate the download links for all projects
            project_links = []
            for project in user_projects:
                print(f"Project ID: {project.id}, Title: {project.title}, File: {project.file}")  # Print basic project info

                # Handle cases where the file field is returned as a list
                file_field = project.file if isinstance(project.file, list) and project.file else []

                if file_field:
                    # Extract the first file name from the list
                    file_name = file_field[0] if len(file_field) > 0 else None

                    if file_name:
                        # Generate the correct file URL
                        #file_url = pb.get_file_url(project, file_name, {})  # Provide an empty dictionary for query_params
                        file_url = pb.get_file_url(project, project.file[0], {})
                        print(f"Generated file URL: {file_url}")  # Debug: Print the file URL
                        project_links.append({'title': project.title, 'file_url': file_url})
                    else:
                        print(f"No valid file found for project: {project.title}")  # Debugging missing file case
                        project_links.append({'title': project.title, 'file_url': None})
                else:
                    print(f"No file found for project: {project.title}")  # Debugging missing file case
                    project_links.append({'title': project.title, 'file_url': None})

            return render_template('download.html', project_links=project_links, new_project=project_title)

        except Exception as e:
            return f"Error fetching user projects: {e}"



    
    return render_template('clause_lookup.html', procurement_types=procurement_types)
    # If GET request, just render the form
    #return render_template('index.html', procurement_types=df_sort1.columns[1:].tolist())


@app.route('/download/<filename>')
def download_file(filename):
    return f"Hit the download route! Requested file: {filename}"
    # Print the filename and the file path for debugging
    file_path = os.path.join(UPLOAD_FOLDER, filename)
    print(f"Requested file: {filename}")
    print(f"Full file path: {file_path}")

    # Check if the file exists
    if os.path.exists(file_path):
        print(f"Serving file: {file_path}")
        # Test with a static download name to see if it works
        return send_file(file_path, as_attachment=True, download_name="test.docx")
    else:
        print("File not found")
        return "File not found", 404

    

if __name__ == '__main__':
    app.run(debug=True)